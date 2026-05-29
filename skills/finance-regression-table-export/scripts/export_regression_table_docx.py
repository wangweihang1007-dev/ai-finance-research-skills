#!/usr/bin/env python
"""Export long-format regression results to a Word three-line table."""

from __future__ import annotations

import argparse
import math
from pathlib import Path
from typing import Iterable

import pandas as pd
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


STAT_TERMS = {
    "n",
    "obs",
    "observations",
    "sample_size",
    "r2",
    "r_squared",
    "adj_r2",
    "adjusted_r2",
    "controls",
    "control_variables",
    "firm_fe",
    "year_fe",
    "industry_fe",
    "province_fe",
}


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description="Export long-format finance regression results to a DOCX three-line table."
    )
    parser.add_argument("--input", required=True, help="Input CSV with term, model, coef, se columns.")
    parser.add_argument("--output", required=True, help="Output .docx path.")
    parser.add_argument("--title", default="表 4-1 回归结果", help="Chinese table title.")
    parser.add_argument(
        "--paren-label",
        default="标准误",
        help="Meaning of parentheses, e.g. 标准误, t值, or 稳健标准误.",
    )
    parser.add_argument("--note", default="", help="Extra note appended below the table.")
    parser.add_argument("--decimals", type=int, default=3, help="Decimal places for numeric values.")
    parser.add_argument("--landscape", action="store_true", help="Use landscape page orientation.")
    parser.add_argument(
        "--no-stars",
        action="store_true",
        help="Do not add significance stars from p-values.",
    )
    return parser.parse_args()


def is_missing(value: object) -> bool:
    if value is None:
        return True
    try:
        return bool(pd.isna(value))
    except TypeError:
        return False


def is_truthy(value: object) -> bool:
    if is_missing(value):
        return False
    return str(value).strip().lower() in {"1", "true", "yes", "y", "stat", "stats"}


def format_number(value: object, decimals: int) -> str:
    if is_missing(value):
        return ""
    if isinstance(value, str):
        stripped = value.strip()
        if stripped == "":
            return ""
        try:
            value = float(stripped)
        except ValueError:
            return stripped
    if isinstance(value, (int, float)):
        if isinstance(value, float) and not math.isfinite(value):
            return ""
        return f"{value:.{decimals}f}"
    return str(value)


def stars_from_p(value: object) -> str:
    if is_missing(value):
        return ""
    try:
        p_value = float(value)
    except (TypeError, ValueError):
        return ""
    if p_value < 0.01:
        return "***"
    if p_value < 0.05:
        return "**"
    if p_value < 0.1:
        return "*"
    return ""


def set_cell_text(cell, text: str, bold: bool = False, align=WD_ALIGN_PARAGRAPH.CENTER) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = align
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.size = Pt(10.5)
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_paragraph_text(paragraph, text: str, size: float, bold: bool = False) -> None:
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")


def set_cell_border(cell, **kwargs) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)

    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        edge_data = kwargs.get(edge)
        if edge_data is None:
            continue
        tag = "w:{}".format(edge)
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        for key, value in edge_data.items():
            element.set(qn(f"w:{key}"), str(value))


def clear_borders(table) -> None:
    nil = {"val": "nil"}
    for row in table.rows:
        for cell in row.cells:
            set_cell_border(cell, top=nil, left=nil, bottom=nil, right=nil, insideH=nil, insideV=nil)


def apply_three_line_borders(table, header_row_idx: int, last_row_idx: int) -> None:
    top_rule = {"val": "single", "sz": "12", "space": "0", "color": "000000"}
    mid_rule = {"val": "single", "sz": "8", "space": "0", "color": "000000"}
    bottom_rule = {"val": "single", "sz": "12", "space": "0", "color": "000000"}

    for cell in table.rows[0].cells:
        set_cell_border(cell, top=top_rule)
    for cell in table.rows[header_row_idx].cells:
        set_cell_border(cell, bottom=mid_rule)
    for cell in table.rows[last_row_idx].cells:
        set_cell_border(cell, bottom=bottom_rule)


def ordered_terms(df: pd.DataFrame) -> list[str]:
    work = df.copy()
    work["_is_stat"] = work.apply(is_stat_row, axis=1)
    terms: list[str] = []
    var_df = work[~work["_is_stat"]].copy()
    if "order" in var_df.columns:
        var_df["_order"] = pd.to_numeric(var_df["order"], errors="coerce")
        var_df = var_df.sort_values(["_order", "term"], na_position="last")
    for term in var_df["term"].dropna().astype(str):
        if term not in terms:
            terms.append(term)

    for term in work[work["_is_stat"]]["term"].dropna().astype(str):
        if term not in terms:
            terms.append(term)
    return terms


def is_stat_row(row: pd.Series) -> bool:
    if "stat" in row.index and is_truthy(row.get("stat")):
        return True
    term = str(row.get("term", "")).strip().lower()
    return term in STAT_TERMS


def label_for(group: pd.DataFrame, term: str) -> str:
    if "label" in group.columns:
        labels = [x for x in group["label"].tolist() if not is_missing(x) and str(x).strip()]
        if labels:
            return str(labels[0]).strip()
    return term


def value_for(group: pd.DataFrame, model: str, column: str) -> object:
    subset = group[group["model"].astype(str) == model]
    if subset.empty or column not in subset.columns:
        return ""
    return subset.iloc[0].get(column, "")


def validate_input(df: pd.DataFrame) -> None:
    required = {"term", "model", "coef", "se"}
    missing = required - set(df.columns)
    if missing:
        raise ValueError(f"Input CSV is missing required columns: {', '.join(sorted(missing))}")


def build_table_rows(df: pd.DataFrame, models: Iterable[str], decimals: int, no_stars: bool) -> list[list[str]]:
    rows: list[list[str]] = []
    for term in ordered_terms(df):
        group = df[df["term"].astype(str) == term]
        stat_row = any(is_stat_row(row) for _, row in group.iterrows())
        label = label_for(group, term)
        coef_values = [label]
        se_values = [""]
        for model in models:
            coef = format_number(value_for(group, model, "coef"), decimals)
            if not stat_row and not no_stars:
                coef = coef + stars_from_p(value_for(group, model, "p"))
            coef_values.append(coef)
            se = format_number(value_for(group, model, "se"), decimals)
            se_values.append(f"({se})" if se and not stat_row else "")
        rows.append(coef_values)
        if not stat_row:
            rows.append(se_values)
    return rows


def export_docx(df: pd.DataFrame, output: Path, title: str, paren_label: str, note: str, decimals: int, landscape: bool, no_stars: bool) -> None:
    validate_input(df)
    output.parent.mkdir(parents=True, exist_ok=True)

    models = list(dict.fromkeys(df["model"].astype(str).tolist()))
    rows = build_table_rows(df, models, decimals, no_stars)

    document = Document()
    section = document.sections[0]
    if landscape:
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width, section.page_height = section.page_height, section.page_width

    title_para = document.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_text(title_para, title, 11, bold=True)

    table = document.add_table(rows=1 + len(rows), cols=1 + len(models))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    header = ["变量"] + models
    for col_idx, value in enumerate(header):
        set_cell_text(table.cell(0, col_idx), value, bold=True)

    for row_idx, row_values in enumerate(rows, start=1):
        for col_idx, value in enumerate(row_values):
            align = WD_ALIGN_PARAGRAPH.LEFT if col_idx == 0 else WD_ALIGN_PARAGRAPH.CENTER
            set_cell_text(table.cell(row_idx, col_idx), value, align=align)

    clear_borders(table)
    apply_three_line_borders(table, header_row_idx=0, last_row_idx=len(rows))

    note_text = f"注：括号内为{paren_label}；***、**、*分别表示在1%、5%、10%的水平上显著。"
    if no_stars:
        note_text = f"注：括号内为{paren_label}。"
    if note:
        note_text = f"{note_text}{note}"
    note_para = document.add_paragraph()
    note_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_paragraph_text(note_para, note_text, 9)

    document.save(output)


def main() -> None:
    args = parse_args()
    input_path = Path(args.input)
    output_path = Path(args.output)
    df = pd.read_csv(input_path)
    export_docx(
        df=df,
        output=output_path,
        title=args.title,
        paren_label=args.paren_label,
        note=args.note,
        decimals=args.decimals,
        landscape=args.landscape,
        no_stars=args.no_stars,
    )
    print(f"Wrote {output_path}")


if __name__ == "__main__":
    main()
